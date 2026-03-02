"""Word 文档解析器"""

from docx import Document


def parse_docx(file_path: str) -> str:
    """
    解析 Word 文档
    
    Args:
        file_path: .docx 文件路径
        
    Returns:
        提取的文本内容
    """
    doc = Document(file_path)
    texts = []

    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                texts.append(row_text)

    return "\n\n".join(texts)
